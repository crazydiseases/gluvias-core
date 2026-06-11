import os
import io
import re
import base64
import logging
import mimetypes
from datetime import datetime
from fastapi import FastAPI, HTTPException, Query, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse, Response, RedirectResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
import httpx
from anthropic import Anthropic
from docx import Document

mimetypes.init()
mimetypes.add_type("text/css", ".css")
mimetypes.add_type("application/javascript", ".js")
mimetypes.add_type("image/svg+xml", ".svg")
mimetypes.add_type("application/json", ".json")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("gluvias-core")

app = FastAPI(title="GLUVIAS // Core System Architecture")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FRONTEND_OUT = os.path.join(BASE_DIR, "frontend", "out")
if not os.path.exists(FRONTEND_OUT):
    FRONTEND_OUT = os.path.abspath("frontend/out")

if os.path.exists(FRONTEND_OUT):
    _next_path = os.path.join(FRONTEND_OUT, "_next")
    if os.path.exists(_next_path):
        app.mount("/_next", StaticFiles(directory=_next_path), name="next-system-assets")

@app.get("/")
async def root_direct():
    return RedirectResponse(url="/dashboard", status_code=307)

COMPANIES_HOUSE_API_URL = "https://api.company-information.service.gov.uk"
RAW_KEY = os.getenv("COMPANIES_HOUSE_KEY", "").strip()
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")

def get_companies_house_headers():
    if not RAW_KEY:
        return {}
    clean_key = RAW_KEY.split(":") if ":" in RAW_KEY else RAW_KEY
    encoded_bytes = base64.b64encode(f"{clean_key}:".encode("utf-8"))
    return {
        "Authorization": f"Basic {encoded_bytes.decode('utf-8')}",
        "Accept": "application/json"
    }

anthropic_client = None
if ANTHROPIC_API_KEY:
    try:
        anthropic_client = Anthropic(api_key=ANTHROPIC_API_KEY)
    except Exception as e:
        logger.error(f"Failed to load Anthropic engine: {str(e)}")

def extract_text_safely(msg_obj) -> str:
    if msg_obj is None:
        return ""
    if hasattr(msg_obj, "content") and isinstance(msg_obj.content, list):
        for block in msg_obj.content:
            if hasattr(block, "text"):
                return str(block.text)
            if isinstance(block, dict) and "text" in block:
                return str(block["text"])
    if hasattr(msg_obj, "text"):
        return str(msg_obj.text)
    if isinstance(msg_obj, dict):
        if "content" in msg_obj and isinstance(msg_obj["content"], list):
            for block in msg_obj["content"]:
                if isinstance(block, dict) and "text" in block:
                    return str(block["text"])
        if "text" in msg_obj:
            return str(msg_obj["text"])
    if isinstance(msg_obj, list) and len(msg_obj) > 0:
        first = msg_obj
        if hasattr(first, "text"):
            return str(first.text)
        if isinstance(first, dict) and "text" in first:
            return str(first["text"])
    return str(msg_obj)

class LegalSearchRequest(BaseModel):
    query: str

@app.post("/api/legal-search")
async def legal_search(req: LegalSearchRequest):
    if not anthropic_client:
        raise HTTPException(status_code=500, detail="Anthropic Key Missing in Environment Config.")
    try:
        msg = anthropic_client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=4000,
            temperature=0.2,
            system="""You are the GLUVIAS Lead Legal Intelligence Architect. Structure your analytical research brief cleanly using these markdown headers exactly:
            ## Executive Legal Summary
            ## Statutory Framework & Legislation Analysis
            ## Relevant Civil Procedure Rules (CPR) Provisions
            ## Leading Jurisprudence & Judicial Precedents
            ## Strategic Litigation Recommendations""",
            messages=[{"role": "user", "content": f"Execute deep legal research on the following parameters:\n\n{req.query}"}]
        )
        return {"intelligence_report": extract_text_safely(msg)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Legal search execution faulted: {str(e)}")


@app.get("/api/company-search")
async def company_search(q: str = Query(..., min_length=1)):
    headers = get_companies_house_headers()
    if not headers:
        raise HTTPException(status_code=401, detail="Authentication token missing or invalid inside environment setup.")
    async with httpx.AsyncClient() as client:
        try:
            res = await client.get(f"{COMPANIES_HOUSE_API_URL}/search/companies", headers=headers, params={"q": q, "items_per_page": 5})
            if res.status_code != 200:
                raise HTTPException(status_code=res.status_code, detail=f"Companies House Registry rejected lookups: {res.text}")
            data = res.json()
            items = data.get("items", []) if isinstance(data, dict) else []
            candidates = []
            for item in items[:5]:
                candidates.append({
                    "name": item.get("title", "Unknown Corporate Body").upper(),
                    "crn": item.get("company_number", ""),
                    "status": item.get("company_status", "Active").upper()
                })
            return {"candidates": candidates}
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Internal Query Processing fault: {str(e)}")

@app.get("/api/company-intelligence")
async def company_intelligence(crn: str = Query(..., min_length=1)):
    headers = get_companies_house_headers()
    if not headers:
        raise HTTPException(status_code=401, detail="Authentication token missing inside environment setup.")
    async with httpx.AsyncClient() as client:
        try:
            p_res = await client.get(f"{COMPANIES_HOUSE_API_URL}/company/{crn}", headers=headers)
            if p_res.status_code != 200:
                raise HTTPException(status_code=p_res.status_code, detail=f"Primary profile failed: {p_res.text}")
                
            # FIX: Pulled request limit threshold up to 100 entries to intercept all partner configurations safely
            o_res = await client.get(f"{COMPANIES_HOUSE_API_URL}/company/{crn}/officers", headers=headers, params={"items_per_page": 100})
            f_res = await client.get(f"{COMPANIES_HOUSE_API_URL}/company/{crn}/filing-history", headers=headers, params={"items_per_page": 20})
            c_res = await client.get(f"{COMPANIES_HOUSE_API_URL}/company/{crn}/charges", headers=headers)
            psc_res = await client.get(f"{COMPANIES_HOUSE_API_URL}/company/{crn}/persons-with-significant-control", headers=headers)
            
            profile = p_res.json()
            officers_data = o_res.json() if o_res.status_code == 200 else {}
            filings_data = f_res.json() if f_res.status_code == 200 else {}
            charges_data = c_res.json() if c_res.status_code == 200 else {}
            psc_data = psc_res.json() if psc_res.status_code == 200 else {}
            
            comp_name = profile.get('company_name', 'Unknown Entity').upper()
            inc_date_str = profile.get('date_of_creation', 'Unknown')
            age_display = "UNKNOWN TIMELINE"
            if inc_date_str != 'Unknown':
                try:
                    inc_date = datetime.strptime(inc_date_str, "%Y-%m-%d")
                    years = datetime.now().year - inc_date.year
                    age_display = f"{years} YEARS OLD ({inc_date.strftime('%d %B %Y').upper()})"
                except:
                    age_display = f"INCORPORATED ON {inc_date_str.upper()}"

            addr_dict = profile.get('registered_office_address', {})
            addr_parts = [addr_dict.get('address_line_1'), addr_dict.get('address_line_2'), addr_dict.get('locality'), addr_dict.get('postal_code')]
            clean_address = ", ".join([p for p in addr_parts if p]).upper() if any(addr_parts) else "NO REGISTRATION ADDRESS CAPTURED"

            officer_lines = []
            # FIX: Removed the restrictive item slice limit to evaluate every single active and resigned partner returned
            for off in officers_data.get("items", []):
                name = off.get("name", "Unknown Officer").upper()
                role = off.get("officer_role", "Director").upper()
                status = "RESIGNED" if off.get("resigned_on") else "ACTIVE"
                appointments_link = off.get("links", {}).get("appointments", "")
                link_count = 1
                if appointments_link:
                    try:
                        app_res = await client.get(f"https://api.company-information.service.gov.uk{appointments_link}", headers=headers)
                        if app_res.status_code == 200:
                            link_count = app_res.json().get("total_count", 1)
                    except:
                        pass
                cross_str = f"Linked to {link_count} corporate allocations" if link_count > 1 else "No alternative directorship footprint"
                officer_lines.append(f"- Officer: {name} | Designation: {role} | Status: {status} | Network Scan: {cross_str}")
            officers_str = "\n".join(officer_lines) if officer_lines else "- No corporate officers registered."
            
            filing_lines = []
            ui_docs_list = []
            for f in filings_data.get("items", []):
                f_date = f.get("date", "N/A")
                f_desc = f.get("description", "System Compliance Log Event")
                f_cat = f.get("category", "General")
                clean_desc = f_desc.replace("-", " ").upper()
                filing_lines.append(f"- Date: {f_date} | Category: {f_cat.upper()} | Update: {clean_desc}")
                
                doc_link = f.get("links", {}).get("document_metadata", "")
                if doc_link and len(ui_docs_list) < 6:
                    doc_id = doc_link.split("/")[-1]
                    download_url = f"https://document-api.company-information.service.gov.uk/document/{doc_id}/content"
                    ui_docs_list.append({
                        "date": f_date,
                        "type": f.get("type", "PDF").upper(),
                        "description": clean_desc,
                        "url": download_url
                    })
            filings_str = "\n".join(filing_lines) if filing_lines else "- No filing history found."

            charges_str = "- No outstanding secured charges or active debt structures recorded."
            c_items = charges_data.get("items", [])
            if c_items:
                charges_str = "\n".join([f"- Instrument: {c.get('particulars', {}).get('description', 'Corporate Charge')} | Status: {c.get('status').upper()} | Registered: {c.get('delivered_on')}" for c in c_items])

            psc_str = "- No individual distribution blocks indexed."
            p_items = psc_data.get("items", [])
            if p_items:
                psc_str = "\n".join([f"- Holder: {p.get('name').upper()} | Allocation Details: {', '.join(p.get('natures_of_control', [])).upper()}" for p in p_items])

            forensic_payload = f"## Overview Narrative Matrix\n- Legal Entity Name: {comp_name}\n- Company Number: {crn}\n- Registry Status: {profile.get('company_status', 'Active').upper()}\n- Age Profile: {age_display}\n- Office Address: {clean_address}\n\n## Corporate Cross-Links & Directorship Mapping\n{officers_str}\n\n## Persons with Significant Control Matrix\n{psc_str}\n\n## Debt Commitments & Corporate Security Registrations\n{charges_str}\n\n## Statutory Compliance Records & Operational Footprint\n{filings_str}"
            report_content = forensic_payload
            
            if anthropic_client:
                try:
                    msg = anthropic_client.messages.create(
                        model="claude-sonnet-4-6",
                        max_tokens=3500,
                        temperature=0.2,
                        system="""You are the Lead Commercial Intelligence and Risk Architect for GLUVIAS. Your job is to convert raw corporate registry data into a balanced, clear executive report.

                        TONE & STYLE DIRECTIVES:
                        1. Provide a professional, objective, and conversational briefing written in standard plain English. Avoid stiff or overly bureaucratic wording.
                        2. Do not use alarmist, paranoid, or sensational terms (e.g., completely avoid phrases like 'CRITICAL EXPOSURE', 'VULNERABILITY', 'SUSPICIOUS NARRATIVE').
                        3. Clearly state the current facts as recorded in the registry. If a directorship is active or an asset charge is outstanding, present it directly without speculation or leaving open-ended questions.

                        CRITICAL SYSTEM PARSING RULES:
                        - Every analytical point or description line MUST start with a hyphen list marker and space (e.g., "- The company maintains a consistent statutory timeline...").
                        - Never format any text as a standalone block paragraph without a leading hyphen.
                        - Match your analysis to these dashboard layout sections exactly:
                        
                        ## Overview Narrative Matrix
                        
                        ## Unexplained Anomalies & Outliers
                        
                        ## Corporate Cross-Links & Directorship Mapping
                        
                        ## Debt Commitments & Corporate Security Registrations
                        
                        ## Statutory Compliance Records & Operational Footprint
                        
                        ## Commercial Counterparty Recommendation Checkpoints""",
                        messages=[{"role": "user", "content": f"Review this comprehensive corporate dataset. Synthesize a conversational, plain English briefing that accurately accounts for all listed officers and partners. Highlight corporate management records and cleanly note if any individuals hold concurrent directorship configurations across multiple entities, assessing potential patterns without speculation or alarmist framing:\n\n{forensic_payload}"}]
                    )
                    report_content = extract_text_safely(msg)
                except Exception as ai_e:
                    raise HTTPException(status_code=502, detail=f"Anthropic Engine Pipeline Error: {str(ai_e)}")

            anomaly_extract = "The registry data displays standard statutory compliance track records with no outlying structural markers."
            match = re.search(r"## Unexplained Anomalies & Outliers\s*\n+([^#]+)", report_content, re.IGNORECASE)
            if match:
                extracted_text = match.group(1).strip()
                clean_text = extracted_text.replace("-", "").strip()
                if len(clean_text) > 10 and "NO MATERIAL" not in clean_text.upper():
                    anomaly_extract = clean_text[:400] + "..."

            return {
                "fact_table": {
                    "name": comp_name,
                    "crn": crn,
                    "status": profile.get('company_status', 'Active').upper(),
                    "age": age_display,
                    "address": clean_address
                },
                "documents": ui_docs_list,
                "intelligence_report": report_content,       
                "detected_anomalies": anomaly_extract        
            }
        except HTTPException as http_e:
            raise http_e
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Registry Data pipeline integration failure: {str(e)}")

class ExportRequest(BaseModel):
    title: str
    content: str

@app.post("/api/export-docx")
async def export_docx(req: ExportRequest):
    doc = Document()
    doc.add_heading(req.title.upper(), level=0)
    for line in req.content.split("\n"):
        if line.strip(): doc.add_paragraph(line.strip())
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return StreamingResponse(
        file_stream, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=GLUVIAS_FORENSIC_DOSSIER_{req.title}.docx"}
    )

@app.get("/dashboard")
async def serve_dashboard():
    for f in ["dashboard/index.html", "dashboard.html"]:
        p = os.path.join(FRONTEND_OUT, f)
        if os.path.exists(p): return FileResponse(p, media_type="text/html")
    raise HTTPException(status_code=404)

@app.get("/{path:path}")
async def catch_all(request: Request, path: str):
    clean = path.strip("/")
    disk = os.path.join(FRONTEND_OUT, clean)
    if os.path.exists(disk) and os.path.isfile(disk):
        return FileResponse(disk, media_type=mimetypes.guess_type(disk))
    html_f = os.path.join(FRONTEND_OUT, f"{clean}.html")
    if os.path.exists(html_f): return FileResponse(html_f, media_type="text/html")
    return Response(content="GLUVIAS SYSTEM FRAMEWORK ACTIVE", media_type="text/plain")
