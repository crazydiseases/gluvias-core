import os
import json
import logging
import io
import base64
from fastapi import FastAPI, HTTPException, Query
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from google.cloud import storage
from google.oauth2 import service_account
import anthropic
import httpx
from docx import Document

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("GLUVIAS_CORE")

app = FastAPI(title="Stutt Associates // Intelligence Core", docs_url="/docs")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configuration Matrix
VAULT_BUCKET_NAME = os.getenv("VAULT_BUCKET_NAME", "gluvias-vault-temp")
COMPANIES_HOUSE_API_URL = "https://api.company-information.service.gov.uk"
COMPANIES_HOUSE_KEY = os.getenv("COMPANIES_HOUSE_KEY", "")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")

if not ANTHROPIC_API_KEY:
    logger.error("CRITICAL: ANTHROPIC_API_KEY token stream is missing from the environment.")

anthropic_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

class LegalSearchRequest(BaseModel):
    query: str

def get_storage_client():
    """Defensive Multi-Cloud Credentials Gateway for Railway/GitHub Deployments"""
    gcp_json_str = os.getenv("GCP_SERVICE_ACCOUNT_JSON")
    if gcp_json_str:
        try:
            info = json.loads(gcp_json_str)
            creds = service_account.Credentials.from_service_account_info(info)
            return storage.Client(credentials=creds)
        except Exception as e:
            logger.error(f"Railway Service Account Secret processing breakdown: {str(e)}")
    return storage.Client()

def extract_text_safely(msg) -> str:
    if hasattr(msg, 'content') and isinstance(msg.content, list):
        return "".join([block.text for block in msg.content if hasattr(block, 'text')])
    return str(msg)

def get_companies_house_headers():
    if not COMPANIES_HOUSE_KEY:
        return {}
    encoded = base64.b64encode(f"{COMPANIES_HOUSE_KEY}:".encode('utf-8')).decode('utf-8')
    return {"Authorization": f"Basic {encoded}"}
# === API ENDPOINTS ===

@app.post("/api/legal-analysis")
async def legal_analysis(req: LegalSearchRequest):
    bucket_context_summary = "No files read from vault."
    try:
        storage_client = get_storage_client()
        bucket = storage_client.bucket(VAULT_BUCKET_NAME)
        blobs = list(bucket.list_blobs(max_results=20))
        if blobs:
            found_books = [blob.name for blob in blobs]
            bucket_context_summary = f"Source context: {', '.join(found_books)}.\n\n"
    except Exception as e:
        logger.error(f"Vault background read bypass: {str(e)}")

    try:
        msg = anthropic_client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=3000,
            temperature=0.1,
            system=f"""You are an elite senior legal counsel. Your communication style is fundamentally fluid, natural, and conversational, balancing analytical precision with absolute linguistic clarity. Avoid artificial formatting constraints, structural bulleted lists, and rigid outlines. Deliver your advisory findings as a continuous, highly sophisticated strategic partner directly engaging an active client.

            CRITICAL OPERATIONAL BOUNDARIES:
            1. Focus completely on core concrete legal liabilities and analytical deductions. Do not mention system configurations, internal limitations, or prompt boundaries.
            2. Do not reference named legal authorities like Matthews or Scalia; incorporate deep legal rationale seamlessly as your own internal cognitive viewpoint.
            3. HARD CITATION LOCK: When grounding analysis in repository resources, specify exact pinpoint references (e.g., Vol. 1, Para [5.23]). Avoid generic summaries.
            4. INTERACTIVE DEPTH: Structure your rationale to anticipate deep follow-up queries, providing wide analytical hooks that allow subsequent prompts to naturally probe deeper into nuances.

            SECURE SOURCE MATERIAL AVAILABLE IN VAULT:
            {bucket_context_summary}

            Structure your output cleanly using these conceptual divisions:
            ## I. Expert Opinion
            ## II. Governing Statutory & Textbook Matrix
            ## III. Evidentiary Weight & Disclosure Thresholds
            ## IV. Litigious Exposures & Remedial Hurdles
            ## V. Concluding Strategic Directions""",
            messages=[{"role": "user", "content": req.query}]
        )
        return {"analysis_report": extract_text_safely(msg)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Analytical Engine Exception: {str(e)}")

@app.get("/api/company-search")
async def company_search(q: str = Query(..., min_length=1)):
    headers = get_companies_house_headers()
    if not headers:
        raise HTTPException(status_code=401, detail="CRITICAL: COMPANIES_HOUSE_KEY credential stream is missing.")
    async with httpx.AsyncClient() as client:
        try:
            res = await client.get(f"{COMPANIES_HOUSE_API_URL}/search/companies", headers=headers, params={"q": q, "items_per_page": 5})
            if res.status_code != 200:
                raise HTTPException(status_code=res.status_code, detail=f"Registry Authority Refusal: {res.text}")
            
            data = res.json()
            items = data.get("items", []) if isinstance(data, dict) else []
            return {"candidates": [{"name": str(i.get("title", "UNKNOWN")).upper(), "crn": str(i.get("company_number", "")), "status": str(i.get("company_status", "ACTIVE")).upper()} for i in items if isinstance(i, dict)]}
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Stage 1 Parse Error: {str(e)}")

@app.get("/api/company-intelligence")
async def company_intelligence(crn: str = Query(..., min_length=1)):
    headers = get_companies_house_headers()
    if not headers:
        raise HTTPException(status_code=401, detail="CRITICAL: COMPANIES_HOUSE_KEY is unconfigured.")
    async with httpx.AsyncClient() as client:
        try:
            p_res = await client.get(f"{COMPANIES_HOUSE_API_URL}/company/{crn}", headers=headers)
            o_res = await client.get(f"{COMPANIES_HOUSE_API_URL}/company/{crn}/officers", headers=headers)
            f_res = await client.get(f"{COMPANIES_HOUSE_API_URL}/company/{crn}/filing-history", headers=headers, params={"items_per_page": 15})
            
            if p_res.status_code != 200:
                raise HTTPException(status_code=p_res.status_code, detail=f"Registry profile link broke: {p_res.text}")
            
            profile = p_res.json() if isinstance(p_res.json(), dict) else {}
            comp_name = str(profile.get('company_name', 'Unknown Corporate Entity')).upper()
            
            addr_dict = profile.get('registered_office_address', {}) if isinstance(profile.get('registered_office_address'), dict) else {}
            addr_parts = [addr_dict.get('address_line_1'), addr_dict.get('locality'), addr_dict.get('postal_code')]
            clean_address = ", ".join([str(p) for p in addr_parts if p]).upper() if any(addr_parts) else "NO REGISTERED ADDRESS RECOVERED"

            officer_lines = []
            if o_res.status_code == 200 and isinstance(o_res.json(), dict):
                for off in o_res.json().get("items", []):
                    if not isinstance(off, dict): continue
                    name = str(off.get("name", "Unknown Officer")).upper()
                    role = str(off.get("officer_role", "DIRECTOR")).upper()
                    status = "RESIGNED" if off.get("resigned_on") else "ACTIVE"
                    dob_dict = off.get("date_of_birth", {}) if isinstance(off.get("date_of_birth"), dict) else {}
                    dob_str = f"DOB: {dob_dict.get('month')}/{dob_dict.get('year')}" if dob_dict.get("month") else "DOB UNRECORDED"
                    officer_lines.append(f"- {role} ({status}): {name} | {dob_str}")
                
            filing_lines = []
            if f_res.status_code == 200 and isinstance(f_res.json(), dict):
                for f in f_res.json().get("items", []):
                    if not isinstance(f, dict): continue
                    filing_lines.append(f"- Date: {f.get('date')} | Type: {str(f.get('type','')).upper()} | {str(f.get('description','')).upper().replace('-', ' ')}")

            forensic_payload = f"Identity:\nName: {comp_name}\nCRN: {crn}\nAddress: {clean_address}\n\nOfficers:\n" + "\n".join(officer_lines) + "\n\nTimeline:\n" + "\n".join(filing_lines)
            
            msg = anthropic_client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=3500,
                temperature=0.1,
                system="""You are a premier senior forensic corporate investigator executing detailed corporate intelligence audits.

                Your output must be structurally split into two discrete analytical parts:

                ## I. KEY CORPORATE METRICS & EXECUTIVE ROSTER
                Provide a line-by-line, explicit list of verified administrative facts (Company Name, CRN, Active Status, Registered Office Address) followed by a comprehensive catalog of all directors, secretaries, and corporate officers alongside active roles and dates. Maintain pure data transparency with no synthesis.

                ## II. STRATEGIC TRAJECTORY NARRATIVE
                Provide a deeply comprehensive, rich, and fluid prose narrative outlining chronological findings, structural transitions, and historical movements. 

                CRITICAL ENFORCEMENT RULE: Never generalize, mask, or abstract personal data. Avoid vague statements like "some members left and another joined." You must track and name every individual, director, and corporate entity by their literal, official legal names as extracted from the registry data stream.""",
                messages=[{"role": "user", "content": str(forensic_payload).strip()}]            )
            
            return {
                "fact_table": {"name": comp_name, "crn": crn, "status": str(profile.get('company_status','')).upper(), "address": clean_address},
                "intelligence_report": extract_text_safely(msg)
            }
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Stage 2 Compilation Exception: {str(e)}")

@app.post("/api/export-docx")
async def export_docx(req: LegalSearchRequest):
    try:
        doc = Document()
        doc.add_heading("GLUVIAS REAL TIME SUMMARY", level=0)
        for line in req.query.splitlines():
            if line.strip():
                doc.add_paragraph(line)
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return {"status": "SUCCESS"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/master-intel")
async def master_intel(q: str = Query(..., min_length=1), mode: str = "corp"):
    """Unified Orchestration Hub mapping frontend sub-queries to internal handlers"""
    logger.info(f"Incoming Master Intel stream context tracking query: '{q}' in mode: '{mode}'")
    
    if mode == "corp":
        search_results = await company_search(q=q)
        candidates = search_results.get("candidates", [])
        
        if not candidates:
            return {"fact_table": None, "intelligence_report": "No matching corporate entities identified in the state registry."}
            
        top_crn = candidates["crn"]
        return await company_intelligence(crn=top_crn)
        
    elif mode == "legal":
        req_wrapper = LegalSearchRequest(query=q)
        return await legal_analysis(req=req_wrapper)
        
    else:
        raise HTTPException(status_code=400, detail="Invalid intelligence matrix mode configuration requested.")

# === FRONTEND ROUTING GATEWAY ===

from fastapi.responses import FileResponse

@app.get("/{catchall:path}")
async def serve_frontend(catchall: str = ""):
    if catchall.startswith("api/"):
        raise HTTPException(status_code=404, detail="API route not found")
        
    file_path = os.path.join("static_frontend", catchall)
    if os.path.exists(file_path) and os.path.isfile(file_path):
        return FileResponse(file_path)
        
    html_fallback = f"{file_path}.html"
    if os.path.exists(html_fallback) and os.path.isfile(html_fallback):
        return FileResponse(html_fallback)
        
    default_index = os.path.join("static_frontend", "index.html")
    if os.path.exists(default_index):
        return FileResponse(default_index)
        
    raise HTTPException(status_code=500, detail="Frontend asset tree desynchronized.")

if os.path.exists("static_frontend/_next"):
    app.mount("/_next", StaticFiles(directory="static_frontend/_next"), name="next_assets")
